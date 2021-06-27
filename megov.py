from bs4 import BeautifulSoup
# from docx import Document
from docx.shared import Pt
from datetime import datetime
import sys, os, codecs

# ファイル情報
pyName = 'MeGov'
pyVer = '0.1'
pyAuthor = 'FTS141'
pyComments = 'https://github.com/fts141/MeGov'


def main():

    # setDocumentProperties()
    # writeHeaderAndFooter()
    writeProvision(soup)

    try:
        exportFile.close()
    except:
        print('<!> ファイルの書き込みに失敗しました。同名のファイルが開かれていませんか？')    


def writeProvision(parent):
    global exportFile

    if parent.find('MainProvision') is not None:
        
        heading = '{}（{}）'.format(lawTitle, lawNum)
        # exportFile.add_heading(heading, level=1)
        exportFile.write('{}\n'.format(headingText(heading, 1)))
        writeProvision(parent.find('MainProvision'))

        for count, supplProvision in enumerate(parent.find_all('SupplProvision')):
            # exportFile.add_page_break()
            heading = '{}'.format(supplProvision.find('SupplProvisionLabel').text)
            if count != 0:
                heading = '{}（{}）'.format(heading ,supplProvision.get('AmendLawNum'))
            exportFile.write('{}\n'.format(headingText(heading, 1)))
            writeProvision(supplProvision)

    elif parent.find('Chapter') is not None:
        for chapter in parent.find_all('Chapter'):
            # exportFile.add_heading(chapter.find('ChapterTitle').text, level=2)
            exportFile.write('{}\n'.format(headingText(chapter.find('ChapterTitle').text, 2)))
            writeProvision(chapter)

    elif parent.find('Section') is not None:
        for section in parent.find_all('Section'):
            # exportFile.add_heading(section.find('SectionTitle').text, level=3)
            exportFile.write('{}\n'.format(headingText(section.find('SectionTitle').text, 3)))
            writeProvision(section)

    elif parent.find('Article') is not None:
        for article in parent.find_all('Article'):
            heading = article.find('ArticleTitle').text
            if article.find('ArticleCaption') is not None:
                heading = '{}{}'.format(heading, article.find('ArticleCaption').text)
            # exportFile.add_heading(heading, level=4)
            exportFile.write('{}\n'.format(headingText(heading, 4)))
            writeProvision(article)

    elif parent.find('Paragraph') is not None:
        for count, paragraph in enumerate(parent.find_all('Paragraph')):
            if count != 0:
                # exportFile.add_heading(paragraph.find('ParagraphNum').text, level=5).paragraph_format.left_indent = indent(1)
                # exportFile.add_paragraph(paragraph.find('ParagraphSentence').find('Sentence').text).paragraph_format.left_indent = indent(1)
                exportFile.write('{}\n'.format(headingText(paragraph.find('ParagraphNum').text, 5)))
            # else:
            #     exportFile.add_paragraph(paragraph.find('ParagraphSentence').find('Sentence').text)
            exportFile.write('{}\n'.format(paragraph.find('ParagraphSentence').find('Sentence').text))
            writeProvision(paragraph)

    elif parent.find('Item') is not None:
        for item in parent.find_all('Item'):
            # exportFile.add_heading(item.find('ItemTitle').text, level=6).paragraph_format.left_indent = indent(2)
            # exportFile.add_paragraph(item.find('ItemSentence').find('Sentence').text).paragraph_format.left_indent = indent(2)
            exportFile.write('{}\n'.format(headingText(item.find('ItemTitle').text, 5)))
            exportFile.write('{}\n'.format(item.find('ItemSentence').find('Sentence').text))

    else:
        pass


# def setDocumentProperties():
#     properties = exportFile.core_properties
#     properties.title = u'{}（{}）'.format(lawTitle, lawNum)
#     properties.author = u'{}'.format(pyAuthor)
#     properties.comments = u'{}'.format(pyComments)
#     properties.language = u'ja-JP'      # 機能していない？


# def writeHeaderAndFooter():
#     global exportFile

#     header = exportFile.sections[0].header.paragraphs[0]
#     header.text = '\t{}\n\t（{}）'.format(lawTitle, lawNum)
#     header.style = exportFile.styles['Header']

#     footer = exportFile.sections[0].footer.paragraphs[0]
#     #footer.text = '\t{} v{} - {}\n\tExport Datetime {}'.format(pyName, pyVer, pyAuthor, datetime.now())
#     footer.text = '\t{} v{} - {}'.format(pyName, pyVer, pyAuthor)
#     footer.style = exportFile.styles['Footer']


def headingText(text, level):
    heading = '#' * level
    return '{} {}'.format(heading, text)


def indentedText(text, level):
    indent = (' ' * 2)* level       # インデントは半角スペース2つ
    return '{}{}'.format(indent, text)


if __name__ == '__main__':

    try:
        soup = BeautifulSoup(open(sys.argv[1], 'r', encoding='UTF-8'), 'lxml-xml')
    except:
        print('''
        <!> ファイルの読み込みに失敗しました。
        実行例 -> python3 megov.py importFile.xml ./exportDir
        ''')
        quit()

    lawTitle = soup.find('LawTitle').text
    lawNum = soup.find('LawNum').text

    try:
        exportDir = '{}/{}/'.format(sys.argv[2], lawTitle)
        if not os.path.exists(exportDir):
            os.makedirs(exportDir)
        exportFileName = '{}（{}）.md'.format(lawTitle, lawNum)
        exportFile = codecs.open('{}{}'.format(exportDir, exportFileName), 'w+', 'utf-8')
    except:
        print('''
        <!> 書き出し先の取得に失敗しました。
        実行例 -> python3 megov.py importFile.xml ./exportDir/
        ''')
        quit()

    main()